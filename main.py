# -*- coding: utf-8 -*-
"""
AI 写作助手 - 智能文稿创作平台
支持 Anthropic Claude、DeepSeek、OpenAI 及自定义兼容接口
支持学术论文、研究报告、工作计划、反思总结、案例分析、工作总结及自定义文稿
"""

import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox
import threading
import json
import os
import re
from datetime import datetime

# ── 引入 docx 相关库用于公文排版 ──────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Markdown 转纯文本工具 ────────────────────────────────────────────────────
def md_to_plain(text: str) -> str:
    """将 Markdown 文本转换为干净的纯文本"""
    text = re.sub(r"```[\s\S]*?```", lambda m: m.group().replace("```", "").strip(), text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"^#{1,6}\s+(.+)$", r"\1", text, flags=re.MULTILINE)
    text = re.sub(r"\*{1,3}([^*]+)\*{1,3}", r"\1", text)
    text = re.sub(r"_{1,3}([^_]+)_{1,3}", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"^>+\s?", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*\d+\.\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ── 公文格式化保存核心逻辑 ────────────────────────────────────────────────────
def save_as_docx(filepath: str, title: str, md_text: str):
    """
    将 Markdown 转换为符合《党政机关公文格式》标准的 Word 文档
    规范参考：GB/T 9704-2012
    """
    
    doc = Document()

    # ── 1. 页面设置 (Page Setup) ──
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(37)
    section.bottom_margin = Mm(35)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)

    # 开启奇偶页页眉页脚不同
    doc.settings.odd_and_even_pages_header_footer = True

    # ── 2. 基础样式定义 (Styles) ──
    def set_run_font(run, font_cn, font_en='Times New Roman', size_pt=16, bold=False):
        run.font.name = font_en
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)

    # 修改默认样式 'Normal' 为公文正文样式
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    style_normal.font.size = Pt(16)
    style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style_normal.paragraph_format.line_spacing = Pt(28)
    style_normal.paragraph_format.first_line_indent = Pt(32)

    # ── 3. 标题排版 (Main Title) ──
    head_p = doc.add_paragraph()
    head_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head_p.paragraph_format.first_line_indent = Pt(0)
    head_p.paragraph_format.line_spacing = Pt(28)
    head_p.paragraph_format.space_before = Pt(0)
    head_p.paragraph_format.space_after = Pt(28) 

    run_title = head_p.add_run(title)
    set_run_font(run_title, '方正小标宋简体', size_pt=22, bold=False)

    # ── 4. 正文内容解析与转换 ──
    lines = md_text.splitlines()
    for line in lines:
        stripped = line.rstrip()
        
        if re.match(r"^[-*_]{3,}\s*$", stripped):
            continue

        # 识别标题 (#)
        heading_match = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = _strip_inline(heading_match.group(2))
            
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = Pt(28)
            p.paragraph_format.first_line_indent = Pt(32)

            run = p.add_run(text)
            
            if level == 1:
                set_run_font(run, 'SimHei', size_pt=16) 
            elif level == 2:
                set_run_font(run, 'KaiTi', size_pt=16)
            else:
                set_run_font(run, '仿宋', size_pt=16, bold=True)
            continue
            
        if not stripped:
            continue

        # 普通段落 (正文)
        p = doc.add_paragraph()
        _add_inline_runs_styled(p, stripped)

    # ── 5. 页码设置 (Page Numbers) ──
    def create_page_number_xml(run):
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run._element.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar2)

    def setup_footer(footer, alignment):
        p = footer.paragraphs[0]
        p.alignment = alignment
        p.paragraph_format.first_line_indent = 0
        r1 = p.add_run("— ") 
        set_run_font(r1, 'SimSun', size_pt=14)
        r2 = p.add_run()
        set_run_font(r2, 'SimSun', size_pt=14)
        create_page_number_xml(r2)
        r3 = p.add_run(" —")
        set_run_font(r3, 'SimSun', size_pt=14)

    setup_footer(section.footer, WD_ALIGN_PARAGRAPH.RIGHT)
    setup_footer(section.even_page_footer, WD_ALIGN_PARAGRAPH.LEFT)

    doc.save(filepath)


def _strip_inline(text: str) -> str:
    """去掉行内 Markdown 符号，只保留文字"""
    text = re.sub(r"\*{1,3}([^*]+)\*{1,3}", r"\1", text)
    text = re.sub(r"_{1,3}([^_]+)_{1,3}", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)
    return text


def _add_inline_runs_styled(paragraph, text: str):
    """解析 Markdown 行内格式并应用到 Docx Run"""
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
    
    pattern = re.compile(r"(\*{1,3}[^*]+\*{1,3}|_{1,3}[^_]+_{1,3}|`[^`]+`)")
    last = 0
    
    def apply_style(run, bold=False, italic=False, code=False):
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0,0,0)
        
        if bold: run.font.bold = True
        if italic: run.font.italic = True
        if code:
             run.font.name = 'Courier New'

    for m in pattern.finditer(text):
        if m.start() > last:
            r = paragraph.add_run(text[last:m.start()])
            apply_style(r)
            
        token = m.group()
        if token.startswith("***") or token.startswith("___"):
            r = paragraph.add_run(token[3:-3])
            apply_style(r, bold=True, italic=True)
        elif token.startswith("**") or token.startswith("__"):
            r = paragraph.add_run(token[2:-2])
            apply_style(r, bold=True)
        elif token.startswith("*") or token.startswith("_"):
            r = paragraph.add_run(token[1:-1])
            apply_style(r, italic=True)
        elif token.startswith("`"):
            r = paragraph.add_run(token[1:-1])
            apply_style(r, code=True)
        last = m.end()
        
    if last < len(text):
        r = paragraph.add_run(text[last:])
        apply_style(r)


# ── 主题配置 ────────────────────────────────────────────────────────────────
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

# ── 常量定义 ────────────────────────────────────────────────────────────────
CONFIG_FILE = os.path.join(os.path.expanduser("~"), ".ai_writer_config.json")
APP_VERSION = "v2.2.3"  # Updated version
APP_AUTHOR  = "Yu JinQuan"

# ── 服务商配置表 ────────────────────────────────────────────────────────────
PROVIDERS = {
    "Anthropic (Claude)": {
        "icon":     "🤖",
        "type":     "anthropic",
        "base_url": "[https://api.anthropic.com](https://api.anthropic.com)",
        "key_hint": "sk-ant-api03-...",
        "models": [
            "claude-3-5-sonnet-20241022",
            "claude-3-opus-20240229",
            "claude-3-haiku-20240307",
        ],
        "default_model": "claude-3-5-sonnet-20241022",
    },
    "DeepSeek": {
        "icon":     "🐋",
        "type":     "openai_compat",
        "base_url": "[https://api.deepseek.com](https://api.deepseek.com)",
        "key_hint": "sk-...",
        "models": [
            "deepseek-chat",
            "deepseek-reasoner",
        ],
        "default_model": "deepseek-chat",
    },
    "OpenAI": {
        "icon":     "🌐",
        "type":     "openai_compat",
        "base_url": "[https://api.openai.com/v1](https://api.openai.com/v1)",
        "key_hint": "sk-...",
        "models": [
            "gpt-4o",
            "gpt-4o-mini",
            "o1-preview",
            "o1-mini",
        ],
        "default_model": "gpt-4o",
    },
    "自定义 (OpenAI 兼容)": {
        "icon":     "🔧",
        "type":     "openai_compat",
        "base_url": "",
        "key_hint": "API Key...",
        "models": [],
        "default_model": "",
    },
}

PROVIDER_NAMES = list(PROVIDERS.keys())

# ── 文稿类型 ────────────────────────────────────────────────────────────────
DOCUMENT_TYPES = [
    ("📄", "学术论文",  "含摘要、引言、方法、结果、讨论、参考文献"),
    ("📊", "研究报告",  "含背景、分析框架、结论与建议"),
    ("📋", "工作计划",  "含目标、阶段步骤、时间线、资源安排"),
    ("🔍", "反思总结",  "含经历回顾、收获、不足与改进方向"),
    ("🔬", "案例分析",  "含案例背景、问题呈现、深度分析、启示"),
    ("📝", "工作总结",  "含工作概述、核心成果、问题与展望"),
    ("✨", "自定义",    "根据您的描述自由定制文稿类型与结构"),
]

# ── 提示词系统 (Prompts) ────────────────────────────────────────────────────
# 为了防止复制截断，已检查字符串完整性
OUTLINE_SYSTEM = (
    "你是一位资深写作顾问，擅长为各类专业文稿设计清晰、合理的结构大纲。\n\n"
    "请根据用户提供的文稿类型、题目和要求，输出一份层次分明的大纲。\n\n"
    "格式规范：\n"
    "- 一级章节：1. 章节名称（简要说明本章核心内容）\n"
    "- 二级章节：1.1 小节名称（说明）\n"
    "- 三级要点：1.1.1 要点（如有必要）\n"
    "- 每个条目要精炼，括号内说明控制在20字以内\n\n"
    "注意：\n"
    "- 直接输出大纲正文，无需前言或解释\n"
    "- 学术论文须包含摘要、关键词、引言、正文各节、结论、参考文献\n"
    "- 其他类型按其行文惯例组织结构\n"
    "- 大纲条目数量适中，一般10~20条为宜"
)

WRITING_SYSTEM = (
    "你是一位经验丰富的专业写作专家，擅长撰写高质量、内容充实的各类文稿。\n\n"
    "请严格依据提供的文稿类型、题目、要求和大纲，撰写完整的正文内容。\n\n"
    "写作规范：\n"
    "- 语言专业、准确、流畅，符合相应文体规范\n"
    "- 内容充实，论据充分，逻辑严密\n"
    "- 严格按照大纲结构依次展开，不得遗漏章节\n"
    "- 每个章节内容饱满，避免空洞\n"
    "- 学术论文须有理论依据，工作类文稿须结合实际\n"
    "- 使用 Markdown 格式：# 一级标题，## 二级标题，**加粗**等\n"
    "- 直接输出正文，无需额外说明"
)


# ── 配置管理器 ──────────────────────────────────────────────────────────────
class ConfigManager:
    def __init__(self):
        self._data = self._load()

    def _default(self):
        import copy
        return {
            "provider":  "Anthropic (Claude)",
            "last_type": "学术论文",
            "providers": {
                pname: {
                    "api_key":  "",
                    "model":    info["default_model"],
                    "base_url": info["base_url"],
                }
                for pname, info in PROVIDERS.items()
            }
        }

    def _load(self):
        try:
            if os.path.exists(CONFIG_FILE):
                with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                    stored = json.load(f)
                # 补全新增服务商
                for pname, info in PROVIDERS.items():
                    stored.setdefault("providers", {})
                    stored["providers"].setdefault(pname, {
                        "api_key":  "",
                        "model":    info["default_model"],
                        "base_url": info["base_url"],
                    })
                return stored
        except Exception:
            pass
        return self._default()

    def save(self):
        try:
            with open(CONFIG_FILE, "w", encoding="utf-8") as f:
                json.dump(self._data, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def get(self, key, default=""):
        return self._data.get(key, default)

    def set(self, key, value):
        self._data[key] = value
        self.save()

    def get_provider_cfg(self, pname):
        return self._data.get("providers", {}).get(pname, {})

    def set_provider_cfg(self, pname, key, value):
        self._data.setdefault("providers", {}).setdefault(pname, {})
        self._data["providers"][pname][key] = value
        self.save()


# ── API 调用层 ──────────────────────────────────────────────────────────────
class APIClient:
    """统一封装 Anthropic 与 OpenAI 兼容接口的流式调用"""

    def __init__(self, provider_name, api_key, model, base_url=""):
        self.provider_name = provider_name
        self.api_key       = api_key
        self.model         = model
        self.base_url      = base_url
        self.ptype         = PROVIDERS[provider_name]["type"]

    def stream(self, system, user_prompt, max_tokens=4096):
        """生成器：逐 token yield 文字片段"""
        if self.ptype == "anthropic":
            yield from self._stream_anthropic(system, user_prompt, max_tokens)
        else:
            yield from self._stream_openai(system, user_prompt, max_tokens)

    def _stream_anthropic(self, system, prompt, max_tokens):
        import anthropic
        client = anthropic.Anthropic(
            api_key=self.api_key,
            base_url=self.base_url if self.base_url else None
        )
        with client.messages.stream(
            model=self.model,
            max_tokens=max_tokens,
            system=system,
            messages=[{"role": "user", "content": prompt}],
        ) as s:
            for chunk in s.text_stream:
                yield chunk

    def _stream_openai(self, system, prompt, max_tokens):
        from openai import OpenAI
        kwargs = {"api_key": self.api_key}
        if self.base_url:
            kwargs["base_url"] = self.base_url
        client = OpenAI(**kwargs)
        stream = client.chat.completions.create(
            model=self.model,
            max_tokens=max_tokens,
            messages=[
                {"role": "system", "content": system},
                {"role": "user",   "content": prompt},
            ],
            stream=True,
        )
        for chunk in stream:
            delta = chunk.choices[0].delta
            if delta and delta.content:
                yield delta.content


# ── 文本编辑器组件 ──────────────────────────────────────────────────────────
class TextEditor(ctk.CTkFrame):
    def __init__(self, parent, font=None, **kwargs):
        super().__init__(parent, fg_color="transparent")
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(0, weight=1)
        _font = font or ctk.CTkFont(size=13)
        self.textbox = ctk.CTkTextbox(self, font=_font, wrap="word", **kwargs)
        self.textbox.grid(row=0, column=0, sticky="nsew")

    def get(self):
        return self.textbox.get("1.0", "end-1c")

    def set(self, text):
        self.textbox.delete("1.0", "end")
        if text:
            self.textbox.insert("1.0", text)

    def append(self, text):
        self.textbox.insert("end", text)
        self.textbox.see("end")

    def clear(self):
        self.textbox.delete("1.0", "end")


# ── 文档类型侧边栏按钮 ──────────────────────────────────────────────────────
class DocTypeButton(ctk.CTkButton):
    ACTIVE_COLOR   = ("#2B6CB0", "#1A4F8A")
    INACTIVE_COLOR = "transparent"
    HOVER_COLOR    = ("#EBF4FF", "#1E3A5F")

    def __init__(self, parent, icon, name, desc, command, **kwargs):
        super().__init__(
            parent, text=f"  {icon}  {name}", anchor="w",
            font=ctk.CTkFont(size=13), height=40, corner_radius=8,
            fg_color=self.INACTIVE_COLOR, hover_color=self.HOVER_COLOR,
            command=command, **kwargs,
        )

    def activate(self):
        self.configure(fg_color=self.ACTIVE_COLOR, font=ctk.CTkFont(size=13, weight="bold"))

    def deactivate(self):
        self.configure(fg_color=self.INACTIVE_COLOR, font=ctk.CTkFont(size=13))


# ── 主应用 ──────────────────────────────────────────────────────────────────
class AIWriterApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self._cfg      = ConfigManager()
        self._busy     = False
        self._doc_type = self._cfg.get("last_type", "学术论文")
        self._type_btns = {}

        self.title(f"✍️  AI 写作助手  {APP_VERSION}  ·  作者：{APP_AUTHOR}")
        self.geometry("1340x840")
        self.minsize(1000, 640)

        self._build_ui()
        self._load_provider_ui()
        self._select_type(self._doc_type, save=False)

    # ── UI 构建 ─────────────────────────────────────────────────────────────
    def _build_ui(self):
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)
        self._build_sidebar()
        self._build_main()

    def _build_sidebar(self):
        sb = ctk.CTkScrollableFrame(
            self, width=260, corner_radius=0,
            fg_color=("#1A2744", "#0F1A33"),
            scrollbar_button_color=("#2A4070", "#1A3060"),
            scrollbar_button_hover_color=("#3A5090", "#2A4070"),
        )
        sb.grid(row=0, column=0, sticky="nsew")
        sb.grid_columnconfigure(0, weight=1)

        # ── Logo ──────────────────────────────────────────────────────────
        logo = ctk.CTkFrame(sb, fg_color="transparent")
        logo.grid(row=0, column=0, sticky="ew", padx=16, pady=(22, 4))
        ctk.CTkLabel(logo, text="✍️", font=ctk.CTkFont(size=28)).pack(side="left")
        col = ctk.CTkFrame(logo, fg_color="transparent")
        col.pack(side="left", padx=(8, 0))
        ctk.CTkLabel(col, text="AI 写作助手",
                     font=ctk.CTkFont(size=16, weight="bold"),
                     text_color="white").pack(anchor="w")
        ctk.CTkLabel(col, text="智能文稿创作平台",
                     font=ctk.CTkFont(size=10),
                     text_color="#7FA8D4").pack(anchor="w")

        ctk.CTkFrame(sb, height=1, fg_color="#2A4070").grid(
            row=1, column=0, sticky="ew", padx=12, pady=8)

        # ── 文稿类型 ──────────────────────────────────────────────────────
        ctk.CTkLabel(sb, text="  文稿类型",
                     font=ctk.CTkFont(size=11, weight="bold"),
                     text_color="#7FA8D4").grid(row=2, column=0, sticky="w", padx=8, pady=(0, 4))

        for idx, (icon, name, desc) in enumerate(DOCUMENT_TYPES):
            btn = DocTypeButton(sb, icon=icon, name=name, desc=desc,
                                command=lambda n=name: self._select_type(n))
            btn.grid(row=3 + idx, column=0, padx=8, pady=2, sticky="ew")
            self._type_btns[name] = btn

        ctk.CTkFrame(sb, height=1, fg_color="#2A4070").grid(
            row=11, column=0, sticky="ew", padx=12, pady=8)

        # ── API 服务商选择 ─────────────────────────────────────────────────
        ctk.CTkLabel(sb, text="  API 服务商",
                     font=ctk.CTkFont(size=11, weight="bold"),
                     text_color="#7FA8D4").grid(row=12, column=0, sticky="w", padx=8, pady=(0, 4))

        self._provider_var = ctk.StringVar(
            value=self._cfg.get("provider", "Anthropic (Claude)"))

        # 用分段按钮展示服务商
        provider_frame = ctk.CTkFrame(sb, fg_color="transparent")
        provider_frame.grid(row=13, column=0, padx=8, pady=(0, 10), sticky="ew")
        provider_frame.grid_columnconfigure((0, 1), weight=1)

        self._provider_btns = {}
        provider_display = [
            ("🤖", "Anthropic (Claude)"),
            ("🐋", "DeepSeek"),
            ("🌐", "OpenAI"),
            ("🔧", "自定义 (OpenAI 兼容)"),
        ]
        for i, (ico, pname) in enumerate(provider_display):
            short = pname.split(" ")[0]
            btn = ctk.CTkButton(
                provider_frame,
                text=f"{ico}\n{short}",
                font=ctk.CTkFont(size=11),
                height=52,
                corner_radius=8,
                fg_color=("#163366", "#0D2244"),
                hover_color=("#1E4A8A", "#152E5C"),
                command=lambda p=pname: self._switch_provider(p),
            )
            btn.grid(row=i // 2, column=i % 2, padx=3, pady=3, sticky="ew")
            self._provider_btns[pname] = btn

        # ── API Key ───────────────────────────────────────────────────────
        self._key_label = ctk.CTkLabel(sb, text="  API Key",
                                        font=ctk.CTkFont(size=11, weight="bold"),
                                        text_color="#7FA8D4")
        self._key_label.grid(row=14, column=0, sticky="w", padx=8, pady=(0, 4))

        self._key_entry = ctk.CTkEntry(
            sb, placeholder_text="sk-...", show="*", height=34,
            fg_color=("#0D1B36", "#0A1228"), border_color="#2A4070",
            text_color="white", placeholder_text_color="#4A6FA0",
        )
        self._key_entry.grid(row=15, column=0, padx=8, pady=(0, 4), sticky="ew")

        self._show_key = False
        self._eye_btn = ctk.CTkButton(
            sb, text="👁  显示 Key", height=28, font=ctk.CTkFont(size=11),
            fg_color="transparent", border_width=1,
            hover_color=("#1E3A5F", "#162D4A"),
            command=self._toggle_key_visibility,
        )
        self._eye_btn.grid(row=16, column=0, padx=8, pady=(0, 10), sticky="ew")

        # ── 模型选择 ──────────────────────────────────────────────────────
        ctk.CTkLabel(sb, text="  模型",
                     font=ctk.CTkFont(size=11, weight="bold"),
                     text_color="#7FA8D4").grid(row=17, column=0, sticky="w", padx=8, pady=(0, 4))

        self._model_var = ctk.StringVar()
        self._model_menu = ctk.CTkOptionMenu(
            sb, variable=self._model_var,
            values=["loading..."],
            height=34, font=ctk.CTkFont(size=12),
            fg_color=("#0D1B36", "#0A1228"),
            button_color=("#2B6CB0", "#1A4F8A"),
            button_hover_color=("#3A82C8", "#2A5FA0"),
        )
        self._model_menu.grid(row=18, column=0, padx=8, pady=(0, 10), sticky="ew")

        # ── Base URL ──────────────────────────────────────────────────────
        self._url_label = ctk.CTkLabel(sb, text="  Base URL (选填/代理)",
                                        font=ctk.CTkFont(size=11, weight="bold"),
                                        text_color="#7FA8D4")
        self._url_entry = ctk.CTkEntry(
            sb, placeholder_text="[https://api.example.com/v1](https://api.example.com/v1)", height=34,
            fg_color=("#0D1B36", "#0A1228"), border_color="#2A4070",
            text_color="white", placeholder_text_color="#4A6FA0",
        )
        self._url_label.grid(row=19, column=0, sticky="w", padx=8, pady=(0, 4))
        self._url_entry.grid(row=20, column=0, padx=8, pady=(0, 10), sticky="ew")

        # ── 自定义模型名（条件显示）──────────────────────────────────────
        self._custom_model_label = ctk.CTkLabel(
            sb, text="  自定义模型名",
            font=ctk.CTkFont(size=11, weight="bold"),
            text_color="#7FA8D4",
        )
        self._custom_model_entry = ctk.CTkEntry(
            sb, placeholder_text="例如：qwen-plus、glm-4...", height=34,
            fg_color=("#0D1B36", "#0A1228"), border_color="#2A4070",
            text_color="white", placeholder_text_color="#4A6FA0",
        )
        self._custom_model_label.grid(row=21, column=0, sticky="w", padx=8, pady=(0, 4))
        self._custom_model_entry.grid(row=22, column=0, padx=8, pady=(0, 10), sticky="ew")

        # ── 保存按钮 ──────────────────────────────────────────────────────
        ctk.CTkButton(
            sb, text="💾  保存设置", height=36,
            font=ctk.CTkFont(size=13, weight="bold"),
            fg_color=("#1A4F8A", "#153D6F"),
            hover_color=("#2B6CB0", "#1A4F8A"),
            command=self._save_settings,
        ).grid(row=23, column=0, padx=8, pady=(4, 12), sticky="ew")

        # ── 作者信息 ──────────────────────────────────────────────────────
        ctk.CTkFrame(sb, height=1, fg_color="#1E2E50").grid(
            row=24, column=0, sticky="ew", padx=12, pady=(0, 8))
        author_frame = ctk.CTkFrame(sb, fg_color="transparent")
        author_frame.grid(row=25, column=0, sticky="ew", padx=12, pady=(0, 20))
        author_frame.grid_columnconfigure(0, weight=1)
        ctk.CTkLabel(
            author_frame,
            text=f"✍️  {APP_VERSION}",
            font=ctk.CTkFont(size=10),
            text_color="#3A5A8A",
        ).grid(row=0, column=0, sticky="w")
        ctk.CTkLabel(
            author_frame,
            text=f"© Author: {APP_AUTHOR}",
            font=ctk.CTkFont(size=10),
            text_color="#3A5A8A",
        ).grid(row=1, column=0, sticky="w")

    # ── 主区域 ──────────────────────────────────────────────────────────────
    def _build_main(self):
        main = ctk.CTkFrame(self, fg_color="transparent")
        main.grid(row=0, column=1, sticky="nsew", padx=(0, 12), pady=12)
        main.grid_columnconfigure(0, weight=1)
        main.grid_rowconfigure(2, weight=1)

        # ── 顶栏 ──────────────────────────────────────────────────────────
        topbar = ctk.CTkFrame(main, fg_color="transparent", height=44)
        topbar.grid(row=0, column=0, sticky="ew", pady=(0, 10))
        topbar.grid_columnconfigure(2, weight=1)
        topbar.grid_propagate(False)

        self._badge = ctk.CTkLabel(
            topbar, text="📄  学术论文",
            font=ctk.CTkFont(size=14, weight="bold"),
            fg_color=("#2B6CB0", "#1A4F8A"),
            corner_radius=8, padx=14, pady=6,
        )
        self._badge.grid(row=0, column=0, padx=(0, 8))

        self._provider_badge = ctk.CTkLabel(
            topbar, text="🤖  Anthropic",
            font=ctk.CTkFont(size=12),
            fg_color=("#163366", "#0D2244"),
            corner_radius=8, padx=10, pady=6,
        )
        self._provider_badge.grid(row=0, column=1, padx=(0, 12))

        self._status_var = tk.StringVar(value="就绪 · 请输入题目后生成大纲")
        ctk.CTkLabel(
            topbar, textvariable=self._status_var,
            font=ctk.CTkFont(size=12), text_color="#7FA8D4",
        ).grid(row=0, column=2, sticky="w")

        # ── 输入区 ────────────────────────────────────────────────────────
        input_card = ctk.CTkFrame(main, corner_radius=10)
        input_card.grid(row=1, column=0, sticky="ew", pady=(0, 10))
        input_card.grid_columnconfigure(1, weight=2)
        input_card.grid_columnconfigure(3, weight=3)

        ctk.CTkLabel(input_card, text="题目 / 主题",
                     font=ctk.CTkFont(size=13, weight="bold"),
                     text_color="#A8C8F0").grid(row=0, column=0, padx=(16, 8), pady=14, sticky="w")
        self._title_entry = ctk.CTkEntry(
            input_card, placeholder_text="输入文稿题目或主题...",
            height=38, font=ctk.CTkFont(size=13),
        )
        self._title_entry.grid(row=0, column=1, padx=(0, 16), pady=14, sticky="ew")

        ctk.CTkLabel(input_card, text="附加要求",
                     font=ctk.CTkFont(size=13, weight="bold"),
                     text_color="#A8C8F0").grid(row=0, column=2, padx=(0, 8), pady=14, sticky="w")
        self._req_entry = ctk.CTkEntry(
            input_card,
            placeholder_text="字数、风格、特定内容要求等（可选）...",
            height=38, font=ctk.CTkFont(size=13),
        )
        self._req_entry.grid(row=0, column=3, padx=(0, 16), pady=14, sticky="ew")

        # ── 标签页 ────────────────────────────────────────────────────────
        self._tabs = ctk.CTkTabview(main, corner_radius=10)
        self._tabs.grid(row=2, column=0, sticky="nsew")
        self._build_outline_tab(self._tabs.add("📋  大纲编辑"))
        self._build_output_tab(self._tabs.add("📄  正文输出"))

        # ── 进度条 ────────────────────────────────────────────────────────
        self._progress = ctk.CTkProgressBar(main, mode="indeterminate", height=4)
        self._progress.grid(row=3, column=0, sticky="ew", pady=(6, 0))
        self._progress.set(0)

    def _build_outline_tab(self, tab):
        tab.grid_columnconfigure(0, weight=1)
        tab.grid_rowconfigure(1, weight=1)

        tb = ctk.CTkFrame(tab, fg_color="transparent")
        tb.grid(row=0, column=0, sticky="ew", pady=(4, 8))

        self._btn_gen_outline = ctk.CTkButton(
            tb, text="🔮  生成大纲",
            font=ctk.CTkFont(size=13, weight="bold"), height=38, width=140,
            command=self._on_gen_outline,
        )
        self._btn_gen_outline.pack(side="left", padx=(0, 8))

        ctk.CTkButton(tb, text="🗑  清空", font=ctk.CTkFont(size=12),
                      height=38, width=72, fg_color="transparent", border_width=1,
                      command=lambda: self._outline_editor.clear()).pack(side="left", padx=(0, 8))

        ctk.CTkButton(
            tb, text="✍️  开始撰写",
            font=ctk.CTkFont(size=13, weight="bold"), height=38, width=140,
            fg_color=("#276749", "#1A4731"), hover_color=("#2F855A", "#22543D"),
            command=self._on_gen_text,
        ).pack(side="left", padx=(0, 12))

        ctk.CTkLabel(tb, text="💡 大纲生成后可直接编辑，修改完成后点击「开始撰写」",
                     font=ctk.CTkFont(size=12), text_color="#7FA8D4").pack(side="left")

        self._outline_editor = TextEditor(
            tab, font=ctk.CTkFont(size=13, family="Consolas"))
        self._outline_editor.grid(row=1, column=0, sticky="nsew")

    def _build_output_tab(self, tab):
        tab.grid_columnconfigure(0, weight=1)
        tab.grid_rowconfigure(1, weight=1)

        tb = ctk.CTkFrame(tab, fg_color="transparent")
        tb.grid(row=0, column=0, sticky="ew", pady=(4, 8))

        self._btn_gen_text = ctk.CTkButton(
            tb, text="✍️  开始撰写",
            font=ctk.CTkFont(size=13, weight="bold"), height=38, width=140,
            fg_color=("#276749", "#1A4731"), hover_color=("#2F855A", "#22543D"),
            command=self._on_gen_text,
        )
        self._btn_gen_text.pack(side="left", padx=(0, 8))

        ctk.CTkButton(tb, text="📋  复制", font=ctk.CTkFont(size=12),
                      height=38, width=72, fg_color="transparent", border_width=1,
                      command=self._copy_output).pack(side="left", padx=(0, 6))

        ctk.CTkButton(tb, text="💾  保存", font=ctk.CTkFont(size=12),
                      height=38, width=72, fg_color="transparent", border_width=1,
                      command=self._save_output).pack(side="left", padx=(0, 12))

        self._wc_var = tk.StringVar(value="字数：0")
        ctk.CTkLabel(tb, textvariable=self._wc_var,
                     font=ctk.CTkFont(size=12), text_color="#7FA8D4").pack(side="left")

        self._output_editor = TextEditor(tab, font=ctk.CTkFont(size=13))
        self._output_editor.grid(row=1, column=0, sticky="nsew")

    # ── 服务商切换逻辑 ───────────────────────────────────────────────────────
    def _switch_provider(self, pname):
        self._provider_var.set(pname)
        self._load_provider_ui()

    def _load_provider_ui(self):
        pname = self._provider_var.get()
        # 容错：若 pname 不在 PROVIDERS 中，回退默认
        if pname not in PROVIDERS:
            pname = "Anthropic (Claude)"
            self._provider_var.set(pname)

        pcfg  = self._cfg.get_provider_cfg(pname)
        pinfo = PROVIDERS[pname]

        # 高亮选中的服务商按钮
        for n, btn in self._provider_btns.items():
            if n == pname:
                btn.configure(fg_color=("#2B6CB0", "#1A4F8A"),
                               font=ctk.CTkFont(size=11, weight="bold"))
            else:
                btn.configure(fg_color=("#163366", "#0D2244"),
                               font=ctk.CTkFont(size=11))

        # Key
        self._key_entry.configure(placeholder_text=pinfo["key_hint"])
        self._key_entry.delete(0, "end")
        self._key_entry.insert(0, pcfg.get("api_key", ""))

        # Base URL Handling
        saved_url = pcfg.get("base_url", "")
        default_url = pinfo.get("base_url", "")
        
        self._url_entry.delete(0, "end")
        self._url_entry.insert(0, saved_url if saved_url else default_url)

        # 重置显示 Key 状态
        if self._show_key:
            self._toggle_key_visibility()

        # 模型 & Custom UI Logic
        is_custom = (pname == "自定义 (OpenAI 兼容)")
        if is_custom:
            self._model_menu.configure(values=["自定义"], state="disabled")
            self._model_var.set("自定义")
            self._custom_model_label.grid()
            self._custom_model_entry.grid()
            self._custom_model_entry.delete(0, "end")
            self._custom_model_entry.insert(0, pcfg.get("model", ""))
        else:
            models = pinfo["models"]
            self._model_menu.configure(values=models, state="normal")
            saved_model = pcfg.get("model", pinfo["default_model"])
            self._model_var.set(saved_model if saved_model in models else models[0])
            self._custom_model_label.grid_remove()
            self._custom_model_entry.grid_remove()

        # 顶栏服务商标签
        icon = pinfo["icon"]
        short = pname.split(" ")[0] if pname != "自定义 (OpenAI 兼容)" else "自定义"
        self._provider_badge.configure(text=f"{icon}  {short}")

    def _toggle_key_visibility(self):
        self._show_key = not self._show_key
        self._key_entry.configure(show="" if self._show_key else "*")
        self._eye_btn.configure(
            text="🔒  隐藏 Key" if self._show_key else "👁  显示 Key")

    def _save_settings(self):
        pname = self._provider_var.get()
        self._cfg.set("provider", pname)
        self._cfg.set_provider_cfg(pname, "api_key", self._key_entry.get().strip())
        self._cfg.set_provider_cfg(pname, "base_url", self._url_entry.get().strip())

        if pname == "自定义 (OpenAI 兼容)":
            self._cfg.set_provider_cfg(pname, "model", self._custom_model_entry.get().strip())
        else:
            self._cfg.set_provider_cfg(pname, "model", self._model_var.get())

        self._set_status("✅  设置已保存")

    # ── 工具方法 ─────────────────────────────────────────────────────────────
    def _select_type(self, name, save=True):
        self._doc_type = name
        for n, btn in self._type_btns.items():
            btn.activate() if n == name else btn.deactivate()
        icon = next((i for i, n, _ in DOCUMENT_TYPES if n == name), "✨")
        self._badge.configure(text=f"{icon}  {name}")
        if save:
            self._cfg.set("last_type", name)

    def _set_status(self, text):
        self._status_var.set(text)

    def _set_busy(self, busy):
        self._busy = busy
        state = "disabled" if busy else "normal"
        self._btn_gen_outline.configure(state=state)
        self._btn_gen_text.configure(state=state)
        if busy:
            self._progress.start()
        else:
            self._progress.stop()
            self._progress.set(0)

    def _build_api_client(self):
        pname = self._provider_var.get()
        if pname not in PROVIDERS:
            pname = "Anthropic (Claude)"
        pcfg = self._cfg.get_provider_cfg(pname)

        key = self._key_entry.get().strip() or pcfg.get("api_key", "")
        if not key:
            messagebox.showerror("缺少 API Key",
                                  f"请为「{pname}」填写 API Key 并保存！")
            return None

        base_url_input = self._url_entry.get().strip()
        base_url = base_url_input if base_url_input else pcfg.get("base_url", "")
        
        is_custom = (pname == "自定义 (OpenAI 兼容)")
        if is_custom:
            model = self._custom_model_entry.get().strip() or pcfg.get("model", "")
            if not base_url:
                messagebox.showerror("缺少 Base URL", "自定义服务商需要填写 Base URL！")
                return None
            if not model:
                messagebox.showerror("缺少模型名", "请填写自定义模型名称！")
                return None
        else:
            model = self._model_var.get()

        return APIClient(
            provider_name=pname,
            api_key=key,
            model=model,
            base_url=base_url,
        )

    def _make_prompt(self, outline=""):
        title = self._title_entry.get().strip()
        req   = self._req_entry.get().strip()
        prompt = f"文稿类型：{self._doc_type}\n题目：{title}"
        if outline:
            prompt += f"\n大纲：\n{outline}"
        if req:
            prompt += f"\n特殊要求：{req}"
        return prompt

    # ── 生成大纲 ─────────────────────────────────────────────────────────────
    def _on_gen_outline(self):
        if self._busy:
            return
        if not self._title_entry.get().strip():
            messagebox.showwarning("提示", "请先输入文稿题目或主题！")
            return
        
        try:
            client = self._build_api_client()
            if not client:
                return
        except Exception as e:
            messagebox.showerror("配置错误", str(e))
            return

        self._set_busy(True)
        self._set_status(f"⏳  [{client.provider_name} · {client.model}]  正在生成大纲...")
        self._outline_editor.clear()
        self._tabs.set("📋  大纲编辑")
        prompt = self._make_prompt()

        def run():
            try:
                for chunk in client.stream(OUTLINE_SYSTEM, prompt, max_tokens=2048):
                    self.after(0, lambda c=chunk: self._outline_editor.append(c))
                self.after(0, lambda: self._set_status(
                    "✅  大纲生成完成 · 可直接编辑后点击「开始撰写」"))
            except Exception as exc:
                self.after(0, lambda e=exc: messagebox.showerror("生成失败", f"连接错误：\n{str(e)}\n\n请检查 API Key 或 Base URL (代理) 设置。"))
                self.after(0, lambda: self._set_status("❌  大纲生成失败"))
            finally:
                self.after(0, lambda: self._set_busy(False))

        threading.Thread(target=run, daemon=True).start()

    # ── 生成正文 ─────────────────────────────────────────────────────────────
    def _on_gen_text(self):
        if self._busy:
            return
        if not self._title_entry.get().strip():
            messagebox.showwarning("提示", "请先输入文稿题目或主题！")
            return
        outline = self._outline_editor.get().strip()
        if not outline:
            messagebox.showwarning("提示", "请先生成或填写大纲内容！")
            return
        
        try:
            client = self._build_api_client()
            if not client:
                return
        except Exception as e:
            messagebox.showerror("配置错误", str(e))
            return

        self._set_busy(True)
        self._set_status(f"⏳  [{client.provider_name} · {client.model}]  正在撰写正文...")
        self._output_editor.clear()
        self._wc_var.set("字数：0")
        self._tabs.set("📄  正文输出")
        prompt = self._make_prompt(outline=outline)

        def run():
            char_count = 0
            try:
                for chunk in client.stream(WRITING_SYSTEM, prompt, max_tokens=8192):
                    char_count += len(chunk)
                    self.after(0, lambda c=chunk: self._output_editor.append(c))
                    self.after(0, lambda n=char_count: self._wc_var.set(f"字数：{n}"))
                self.after(0, lambda: self._set_status(
                    f"✅  撰写完成 · [{client.provider_name} · {client.model}] · 共 {char_count} 字"))
            except Exception as exc:
                self.after(0, lambda e=exc: messagebox.showerror("生成失败", f"连接错误：\n{str(e)}\n\n请检查 API Key 或 Base URL (代理) 设置。"))
                self.after(0, lambda: self._set_status("❌  撰写失败"))
            finally:
                self.after(0, lambda: self._set_busy(False))

        threading.Thread(target=run, daemon=True).start()

    # ── 复制 / 保存 ──────────────────────────────────────────────────────────
    def _copy_output(self):
        text = self._output_editor.get()
        if not text:
            messagebox.showinfo("提示", "暂无可复制的内容。")
            return
        self.clipboard_clear()
        self.clipboard_append(text)
        self._set_status("✅  已复制到剪贴板")

    def _save_output(self):
        text = self._output_editor.get()
        if not text:
            messagebox.showinfo("提示", "暂无可保存的内容。")
            return
        fmt = self._ask_save_format()
        if fmt is None:
            return

        title = self._title_entry.get().strip() or "文稿"
        ts    = datetime.now().strftime("%Y%m%d_%H%M%S")

        fmt_cfg = {
            "docx": (".docx", "Word 文档 (公文版式) (*.docx)", "*.docx"),
            "txt":  (".txt",  "纯文本 (*.txt)",             "*.txt"),
            "md":   (".md",   "Markdown (*.md)",            "*.md"),
        }
        def_ext, ftype_name, ftype_glob = fmt_cfg[fmt]

        fp = filedialog.asksaveasfilename(
            defaultextension=def_ext,
            filetypes=[(ftype_name, ftype_glob), ("所有文件", "*.*")],
            initialfile=f"{title}_{ts}",
            title="保存文稿",
        )
        if not fp:
            return

        if not fp.lower().endswith(def_ext):
            fp += def_ext

        try:
            if fmt == "docx":
                save_as_docx(fp, title, text)
            elif fmt == "txt":
                with open(fp, "w", encoding="utf-8") as f:
                    f.write(md_to_plain(text))
            else:
                with open(fp, "w", encoding="utf-8") as f:
                    f.write(text)
            self._set_status(f"✅  已保存：{os.path.basename(fp)}")
        except ImportError:
            messagebox.showerror(
                "缺少依赖",
                "保存 Word 文档需要安装 python-docx：\n\npip install python-docx"
            )
        except Exception as exc:
            import traceback
            traceback.print_exc()
            messagebox.showerror("保存失败", str(exc))

    def _ask_save_format(self):
        """弹出格式选择窗口"""
        result = [None]
        BG       = "#1A2744"
        FG       = "#E8F0FE"
        BTN_BG   = "#163366"
        BTN_HV   = "#2B6CB0"
        CANCEL   = "#0F1A33"
        BORDER   = "#2A4070"

        self.update_idletasks()
        W, H = 320, 230
        x = self.winfo_x() + (self.winfo_width()  - W) // 2
        y = self.winfo_y() + (self.winfo_height() - H) // 2

        dlg = tk.Toplevel(self)
        dlg.title("选择保存格式")
        dlg.geometry(f"{W}x{H}+{x}+{y}")
        dlg.resizable(False, False)
        dlg.configure(bg=BG)
        dlg.transient(self)
        dlg.lift()
        dlg.update()
        dlg.grab_set()
        dlg.focus_force()

        tk.Label(
            dlg, text="请选择保存格式",
            bg=BG, fg=FG,
            font=("TkDefaultFont", 13, "bold"),
        ).pack(pady=(18, 10))

        formats = [
            ("docx", "📝  Word 文档 (公文版式)"),
            ("txt",  "📄  纯文本      (.txt)"),
            ("md",   "🔖  Markdown   (.md)"),
        ]
        for fmt, label in formats:
            btn = tk.Button(
                dlg, text=label,
                bg=BTN_BG, fg=FG, activebackground=BTN_HV, activeforeground=FG,
                relief="flat", bd=0, pady=6,
                font=("TkDefaultFont", 12),
                cursor="hand2",
                command=lambda f=fmt: (result.__setitem__(0, f), dlg.destroy()),
            )
            btn.pack(fill="x", padx=28, pady=3)

        tk.Frame(dlg, bg=BORDER, height=1).pack(fill="x", padx=28, pady=(8, 0))

        tk.Button(
            dlg, text="取消",
            bg=CANCEL, fg="#7FA8D4", activebackground="#1A2744",
            relief="flat", bd=0, pady=5,
            font=("TkDefaultFont", 11),
            cursor="hand2",
            command=dlg.destroy,
        ).pack(fill="x", padx=28, pady=(4, 0))

        dlg.wait_window()
        return result[0]


# ── 入口 ────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    app = AIWriterApp()
    app.mainloop()
