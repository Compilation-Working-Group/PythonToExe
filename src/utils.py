"""
实用工具模块
"""

import re
import json
from datetime import datetime
from typing import List, Dict, Any
import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import html2text

class DocumentFormatter:
    """文档格式化器"""
    
    @staticmethod
    def format_markdown_to_text(markdown_text: str) -> str:
        """将Markdown格式化为纯文本"""
        # 移除Markdown标记
        text = re.sub(r'#{1,6}\s*', '', markdown_text)  # 移除标题标记
        text = re.sub(r'\*{1,2}(.*?)\*{1,2}', r'\1', text)  # 移除粗体和斜体
        text = re.sub(r'!\[.*?\]\(.*?\)', '', text)  # 移除图片
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # 移除链接
        text = re.sub(r'`{1,3}(.*?)`{1,3}', r'\1', text)  # 移除代码标记
        text = re.sub(r'^\s*[-*+]\s*', '', text, flags=re.MULTILINE)  # 移除列表标记
        text = re.sub(r'^\s*\d+\.\s*', '', text, flags=re.MULTILINE)  # 移除数字列表
        
        return text.strip()
    
    @staticmethod
    def format_text_for_docx(text: str) -> List[Dict[str, Any]]:
        """将文本格式化为Word文档结构"""
        sections = []
        lines = text.split('\n')
        
        current_section = {"level": 0, "title": "", "content": []}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 检测标题
            match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if match:
                if current_section["title"]:
                    sections.append(current_section.copy())
                
                level = len(match.group(1))
                title = match.group(2)
                current_section = {"level": level, "title": title, "content": []}
            else:
                current_section["content"].append(line)
        
        if current_section["title"]:
            sections.append(current_section)
        
        return sections
    
    @staticmethod
    def extract_keywords(text: str, num_keywords: int = 5) -> List[str]:
        """从文本中提取关键词"""
        # 简单的关键词提取逻辑
        words = re.findall(r'\b[\u4e00-\u9fff]{2,5}\b', text)
        word_freq = {}
        
        for word in words:
            if len(word) >= 2:
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # 按频率排序
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        
        return [word for word, _ in sorted_words[:num_keywords]]

class FileHandler:
    """文件处理器"""
    
    @staticmethod
    def save_document(content: str, filename: str, format_type: str = "txt"):
        """
        保存文档
        
        Args:
            content: 文档内容
            filename: 文件名
            format_type: 文件格式
        """
        try:
            if format_type == "docx":
                FileHandler.save_as_docx(content, filename)
            elif format_type == "pdf":
                FileHandler.save_as_pdf(content, filename)
            elif format_type == "md":
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(content)
            else:  # txt
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(content)
            return True
        except Exception as e:
            print(f"保存文件失败: {e}")
            return False
    
    @staticmethod
    def save_as_docx(content: str, filename: str):
        """保存为Word文档"""
        doc = Document()
        
        # 设置默认样式
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        
        sections = DocumentFormatter.format_text_for_docx(content)
        
        for section in sections:
            level = section["level"]
            title = section["title"]
            content_lines = section["content"]
            
            # 添加标题
            if level == 1:
                heading = doc.add_heading(title, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                doc.add_heading(title, 1)
            elif level == 3:
                doc.add_heading(title, 2)
            else:
                doc.add_heading(title, 3)
            
            # 添加内容
            for line in content_lines:
                if line.strip():
                    para = doc.add_paragraph(line)
        
        doc.save(filename)
    
    @staticmethod
    def save_as_pdf(content: str, filename: str):
        """保存为PDF文件"""
        # 需要安装reportlab或其他PDF库
        # 这里简化处理，先保存为文本
        with open(filename.replace('.pdf', '.txt'), 'w', encoding='utf-8') as f:
            f.write(content)
        print("PDF保存功能需要安装额外的库，已保存为文本文件")

class TextProcessor:
    """文本处理器"""
    
    @staticmethod
    def estimate_word_count(text: str) -> int:
        """估算字数"""
        # 中文字数估算
        chinese_chars = re.findall(r'[\u4e00-\u9fff]', text)
        english_words = re.findall(r'\b[a-zA-Z]+\b', text)
        
        return len(chinese_chars) + len(english_words)
    
    @staticmethod
    def extract_sections(text: str) -> Dict[str, str]:
        """提取文档各部分"""
        sections = {}
        current_section = "引言"
        current_content = []
        
        lines = text.split('\n')
        
        for line in lines:
            # 检测章节标题
            section_match = re.match(r'^#+\s*(.*?)\s*$', line)
            if section_match and len(section_match.group(1).strip()) < 50:
                if current_section and current_content:
                    sections[current_section] = '\n'.join(current_content).strip()
                
                current_section = section_match.group(1).strip()
                current_content = []
            else:
                if line.strip():
                    current_content.append(line)
        
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        return sections
    
    @staticmethod
    def add_page_numbers(text: str) -> str:
        """添加页码标记"""
        lines = text.split('\n')
        processed_lines = []
        
        for i, line in enumerate(lines, 1):
            if i % 40 == 0:  # 每40行添加一个页码
                page_num = i // 40 + 1
                processed_lines.append(f"\n--- 第{page_num}页 ---\n")
            processed_lines.append(line)
        
        return '\n'.join(processed_lines)
